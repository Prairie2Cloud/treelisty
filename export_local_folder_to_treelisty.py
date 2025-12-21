"""
TreeListy Local Folder Exporter with RAG Content Extraction
Scans a local folder and creates a TreeListy-compatible JSON file using the filesystem pattern.
Extracts text content from documents for RAG (retrieval-augmented generation).

No API setup required - just point to a local folder (like Google Drive sync folder).

Usage:
  python export_local_folder_to_treelisty.py B:/ai_boneyard --extract-content
  python export_local_folder_to_treelisty.py "B:/My Drive/Projects" --max-depth 5

Options:
  --extract-content    Extract text from documents (PDFs, Docs, etc.)
  --max-depth N        Maximum folder depth (default: 10)
  --max-content-size N Max content size per file in KB (default: 100)

Install extraction libraries:
  pip install PyPDF2 python-docx
"""

import os
import sys
import json
import argparse
from datetime import datetime
from pathlib import Path

# Fix Windows console encoding for emojis
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Optional imports for content extraction
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# =============================================================================
# Configuration
# =============================================================================

# File extensions that support content extraction
EXTRACTABLE_EXTENSIONS = {
    '.pdf': 'PDF',
    '.docx': 'Word Doc',
    '.txt': 'Text',
    '.md': 'Markdown',
    '.csv': 'CSV',
    '.json': 'JSON',
}

# Stats tracking
extraction_stats = {
    'attempted': 0,
    'succeeded': 0,
    'failed': 0,
    'skipped': 0,
    'total_chars': 0
}


def get_file_icon(path: Path):
    """Get emoji icon based on file type"""
    if path.is_dir():
        name_lower = path.name.lower()
        if 'download' in name_lower:
            return '📥'
        elif 'document' in name_lower or 'docs' in name_lower:
            return '📄'
        elif 'picture' in name_lower or 'photo' in name_lower or 'image' in name_lower:
            return '🖼️'
        elif 'video' in name_lower or 'movie' in name_lower:
            return '🎬'
        elif 'music' in name_lower or 'audio' in name_lower:
            return '🎵'
        return '📁'

    ext = path.suffix.lower()
    ext_icons = {
        '.pdf': '📕',
        '.doc': '📘', '.docx': '📘',
        '.xls': '📗', '.xlsx': '📗', '.csv': '📊',
        '.ppt': '📙', '.pptx': '📙',
        '.txt': '📝', '.md': '📝',
        '.jpg': '🖼️', '.jpeg': '🖼️', '.png': '🖼️', '.gif': '🎨',
        '.mp4': '🎬', '.avi': '🎬', '.mov': '🎬',
        '.mp3': '🎵', '.wav': '🎵', '.flac': '🎵',
        '.zip': '🗜️', '.rar': '📦', '.7z': '📦',
        '.html': '🌐', '.css': '🎨', '.js': '⚡', '.py': '🐍',
    }
    return ext_icons.get(ext, '📄')


# =============================================================================
# Content Extraction
# =============================================================================

def extract_text_from_pdf(file_path: Path):
    """Extract text from PDF file"""
    if not HAS_PYPDF2:
        return None, "PyPDF2 not installed"

    try:
        reader = PyPDF2.PdfReader(str(file_path))
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            except Exception as e:
                text_parts.append(f"[Page {page_num + 1}] (extraction error: {e})")

        return '\n\n'.join(text_parts), None
    except Exception as e:
        return None, str(e)


def extract_text_from_docx(file_path: Path):
    """Extract text from DOCX file"""
    if not HAS_DOCX:
        return None, "python-docx not installed"

    try:
        doc = docx.Document(str(file_path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return '\n\n'.join(text_parts), None
    except Exception as e:
        return None, str(e)


def extract_text_from_plain(file_path: Path):
    """Extract text from plain text files"""
    try:
        # Try UTF-8 first, then fall back to latin-1
        try:
            return file_path.read_text(encoding='utf-8'), None
        except UnicodeDecodeError:
            return file_path.read_text(encoding='latin-1'), None
    except Exception as e:
        return None, str(e)


def extract_content(file_path: Path, max_size_kb=100):
    """
    Extract text content from a file.

    Args:
        file_path: Path to the file
        max_size_kb: Maximum content size in KB

    Returns:
        tuple: (extracted_text, error_message)
    """
    global extraction_stats
    extraction_stats['attempted'] += 1

    ext = file_path.suffix.lower()
    if ext not in EXTRACTABLE_EXTENSIONS:
        extraction_stats['skipped'] += 1
        return None, "Unsupported type"

    # Check file size
    try:
        file_size = file_path.stat().st_size
        if file_size > max_size_kb * 1024 * 10:  # 10x limit for raw files
            extraction_stats['skipped'] += 1
            return None, f"File too large ({file_size // 1024}KB)"
    except Exception:
        pass

    try:
        # Extract based on type
        if ext == '.pdf':
            text, error = extract_text_from_pdf(file_path)
        elif ext == '.docx':
            text, error = extract_text_from_docx(file_path)
        else:
            # Plain text, markdown, CSV, JSON
            text, error = extract_text_from_plain(file_path)

        if error:
            extraction_stats['failed'] += 1
            return None, error

        # Truncate if too large
        max_chars = max_size_kb * 1024
        if text and len(text) > max_chars:
            text = text[:max_chars] + f"\n\n[Truncated at {max_size_kb}KB]"

        extraction_stats['succeeded'] += 1
        extraction_stats['total_chars'] += len(text) if text else 0
        return text, None

    except Exception as e:
        extraction_stats['failed'] += 1
        return None, str(e)


# =============================================================================
# Folder Scanning
# =============================================================================

def scan_folder(folder_path: Path, depth=0, max_depth=10, extract=False, max_content_kb=100):
    """
    Recursively scan local folder

    Args:
        folder_path: Path to folder to scan
        depth: Current recursion depth
        max_depth: Maximum depth to scan
        extract: Whether to extract content from documents
        max_content_kb: Maximum content size per file in KB

    Returns:
        list: Children nodes in TreeListy filesystem format
    """
    if depth > max_depth:
        print(f"⚠️  Max depth {max_depth} reached at: {folder_path.name}")
        return []

    indent = '  ' * depth
    print(f"{indent}📂 {folder_path.name}")

    try:
        children = []
        items = sorted(folder_path.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower()))

        for item in items:
            # Skip hidden files and system files
            if item.name.startswith('.') or item.name.startswith('~'):
                continue

            is_folder = item.is_dir()
            icon = get_file_icon(item)

            # Get file stats
            try:
                stats = item.stat()
                file_size = stats.st_size if not is_folder else 0
                modified_time = datetime.fromtimestamp(stats.st_mtime).isoformat()
                created_time = datetime.fromtimestamp(stats.st_ctime).isoformat()
            except Exception:
                file_size = 0
                modified_time = ''
                created_time = ''

            # Build node
            node = {
                'id': str(item.resolve()),  # Use full path as ID
                'name': item.name,
                'type': 'item',
                'icon': icon,
                'isFolder': is_folder,
                'fileExtension': item.suffix if not is_folder else '',
                'fileSize': file_size,
                'dateModified': modified_time,
                'dateCreated': created_time,
                'filePath': str(item.resolve())
            }

            # Extract content if enabled and file type is supported
            if extract and not is_folder and item.suffix.lower() in EXTRACTABLE_EXTENSIONS:
                type_label = EXTRACTABLE_EXTENSIONS[item.suffix.lower()]
                print(f"{indent}  📄 Extracting: {item.name} ({type_label})")

                content, error = extract_content(item, max_content_kb)

                if content:
                    # Store content in description for RAG
                    node['description'] = content

                    # Add RAG metadata
                    node['_rag'] = {
                        'source': item.name,
                        'sourceType': 'local-file',
                        'sourceId': str(item.resolve()),
                        'sourcePath': str(item.resolve()),
                        'imported': datetime.now().isoformat(),
                        'charCount': len(content)
                    }
                    print(f"{indent}    ✓ {len(content):,} chars")
                elif error:
                    print(f"{indent}    ⚠️ {error}")

            # Recursively scan subfolders
            if is_folder:
                subchildren = scan_folder(
                    item, depth + 1, max_depth, extract, max_content_kb
                )
                if subchildren:
                    node['children'] = subchildren
                node['expanded'] = False  # Collapsed by default

            children.append(node)

        print(f"{indent}  ✓ {len(children)} items")
        return children

    except PermissionError:
        print(f"{indent}  ❌ Permission denied")
        return []
    except Exception as e:
        print(f"{indent}  ❌ Error: {e}")
        return []


def count_items(children):
    """Count total items recursively"""
    if not children:
        return 0
    count = len(children)
    for child in children:
        if 'children' in child:
            count += count_items(child['children'])
    return count


# =============================================================================
# Main Export
# =============================================================================

def export_folder(folder_path: Path, max_depth=10, extract=False, max_content_kb=100):
    """Main export function"""
    print("\n🌳 TreeListy Local Folder Exporter")
    print("=" * 60)
    print(f"Target folder: {folder_path}")
    print(f"Max scan depth: {max_depth} levels")
    print(f"Content extraction: {'✅ Enabled' if extract else '❌ Disabled'}")
    if extract:
        print(f"Max content size: {max_content_kb}KB per file")
        if not HAS_PYPDF2:
            print("⚠️  PyPDF2 not installed - PDF extraction disabled")
        if not HAS_DOCX:
            print("⚠️  python-docx not installed - DOCX extraction disabled")
    print()

    # Scan folder
    print("📥 Scanning folder...\n")
    children = scan_folder(folder_path, 0, max_depth, extract, max_content_kb)

    # Create TreeListy structure (filesystem pattern) with source metadata
    tree = {
        'id': 'root-local',
        'name': '💻 My Computer',
        'type': 'root',
        'icon': '💻',
        'expanded': True,
        'source': {
            'type': 'local-folder',
            'folderPath': str(folder_path.resolve()),
            'folderName': folder_path.name,
            'lastSync': datetime.now().isoformat(),
            'syncDepth': max_depth,
            'contentExtracted': extract
        },
        'children': [
            {
                'id': str(folder_path.resolve()),
                'name': f'📁 {folder_path.name}',
                'type': 'phase',
                'icon': '📁',
                'expanded': True,
                'children': children
            }
        ],
        'pattern': {
            'key': 'filesystem',
            'labels': None
        }
    }

    # Save to file
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_name = folder_path.name.replace('/', '-').replace('\\', '-').replace(' ', '-')
    output_file = f'local-folder-{safe_name}-{timestamp}.json'

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(tree, f, indent=2, ensure_ascii=False)

    # Summary
    total_items = count_items(children)
    print("\n" + "=" * 60)
    print(f"✅ SUCCESS! Exported to: {output_file}")
    print(f"\n📊 Statistics:")
    print(f"   Total items: {total_items}")
    print(f"   Top-level items: {len(children)}")

    if extract:
        print(f"\n📝 Content Extraction:")
        print(f"   Attempted: {extraction_stats['attempted']}")
        print(f"   Succeeded: {extraction_stats['succeeded']}")
        print(f"   Failed: {extraction_stats['failed']}")
        print(f"   Skipped: {extraction_stats['skipped']}")
        print(f"   Total characters: {extraction_stats['total_chars']:,}")

    print(f"\n📋 Next Steps:")
    print(f"   1. Open TreeListy in browser")
    print(f"   2. Click '📂 Import' → Select '{output_file}'")
    print(f"   3. Use retrieve_context via MCP to search content!")
    print("=" * 60)

    return output_file


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Export local folder to TreeListy with optional content extraction',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python export_local_folder_to_treelisty.py B:/ai_boneyard
      Export folder (metadata only)

  python export_local_folder_to_treelisty.py B:/ai_boneyard --extract-content
      Export folder with text extraction

  python export_local_folder_to_treelisty.py "B:/My Drive/Projects" --max-depth 5
      Export with limited depth

  python export_local_folder_to_treelisty.py . --extract-content
      Export current directory with content
        """
    )

    parser.add_argument('folder', type=str,
                        help='Path to folder to export')
    parser.add_argument('--extract-content', action='store_true',
                        help='Extract text content from documents (PDFs, Docs, etc.)')
    parser.add_argument('--max-depth', type=int, default=10,
                        help='Maximum folder depth (default: 10)')
    parser.add_argument('--max-content-size', type=int, default=100,
                        help='Maximum content size per file in KB (default: 100)')

    args = parser.parse_args()

    folder_path = Path(args.folder).resolve()

    if not folder_path.exists():
        print(f"❌ Folder not found: {folder_path}")
        sys.exit(1)

    if not folder_path.is_dir():
        print(f"❌ Not a directory: {folder_path}")
        sys.exit(1)

    # Run export
    export_folder(
        folder_path=folder_path,
        max_depth=args.max_depth,
        extract=args.extract_content,
        max_content_kb=args.max_content_size
    )


if __name__ == '__main__':
    main()
